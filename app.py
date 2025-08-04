import streamlit as st
import pandas as pd
import PyPDF2
import re
from io import BytesIO
from docx import Document
from openpyxl import Workbook
import matplotlib.pyplot as plt

# --- GROQ Placeholder Function (Replace with actual API integration) ---
def groq_analyze(text, industry):
    return {
        "benchmarks": {
            "Revenue": 5000000,
            "Gross Profit": 2000000,
            "EBITDA": 1000000,
            "Net Income": 500000,
            "Current Ratio": 1.5,
            "Debt-to-Equity": 0.8,
            "ROE": 12.0,
            "ROA": 7.0,
            "Operating Cash Flow": 600000,
            "Free Cash Flow": 400000
        },
        "analysis": f"Compared to the {industry} industry, the company outperforms in EBITDA and Net Income but lags in Free Cash Flow.",
        "recommendations": "Improve working capital management and explore opportunities to reduce debt for better financial leverage."
    }

# --- KPI Extraction ---
def extract_kpis(text):
    kpis = {}
    patterns = {
        "Revenue": r"Revenue[^0-9]*([\d,\.]+)",
        "Gross Profit": r"Gross\s+Profit[^0-9]*([\d,\.]+)",
        "EBITDA": r"EBITDA[^0-9]*([\d,\.]+)",
        "Net Income": r"Net\s+Income[^0-9]*([\d,\.]+)",
        "Current Ratio": r"Current\s+Ratio[^0-9]*([\d,\.]+)",
        "Debt-to-Equity": r"Debt[-\s]*to[-\s]*Equity[^0-9]*([\d,\.]+)",
        "ROE": r"ROE[^0-9]*([\d,\.]+)",
        "ROA": r"ROA[^0-9]*([\d,\.]+)",
        "Operating Cash Flow": r"Operating\s+Cash\s+Flow[^0-9]*([\d,\.]+)",
        "Free Cash Flow": r"Free\s+Cash\s+Flow[^0-9]*([\d,\.]+)"
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                kpis[key] = float(match.group(1).replace(',', ''))
            except ValueError:
                continue
    return kpis

# --- Excel Report ---
def generate_excel(kpis, benchmarks):
    wb = Workbook()
    ws = wb.active
    ws.title = "KPI Comparison"
    ws.append(["Metric", "Company", "Benchmark", "Variance", "Variance %"])
    for k, v in kpis.items():
        benchmark = benchmarks.get(k, 0)
        variance = v - benchmark
        var_pct = (variance / benchmark * 100) if benchmark else 0
        ws.append([k, v, benchmark, variance, var_pct])
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# --- Word Report ---
def generate_word(analysis, recommendations):
    doc = Document()
    doc.add_heading("AI-Generated Financial Analysis", level=1)
    doc.add_paragraph("🧠 Analysis:\n" + analysis)
    doc.add_paragraph("\n💡 Recommendations:\n" + recommendations)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- PDF Report (Bar Chart) ---
def generate_pdf(kpis, benchmarks):
    fig, ax = plt.subplots(figsize=(10, 6))
    labels = list(kpis.keys())
    values = [kpis[k] for k in labels]
    benchmark_values = [benchmarks.get(k, 0) for k in labels]
    x = range(len(labels))
    ax.bar(x, values, width=0.4, label="Company", color='steelblue')
    ax.bar([i + 0.4 for i in x], benchmark_values, width=0.4, label="Benchmark", color='lightgray')
    ax.set_xticks([i + 0.2 for i in x])
    ax.set_xticklabels(labels, rotation=45, ha='right')
    ax.set_ylabel("Amount")
    ax.set_title("KPI vs Industry Benchmark")
    ax.legend()
    fig.tight_layout()
    output = BytesIO()
    fig.savefig(output, format='pdf')
    output.seek(0)
    return output

# --- Streamlit App ---
def main():
    st.set_page_config("Financial Statement Analyzer", layout="wide")
    st.title("📊 Financial Statement Analyzer")
    st.markdown("Upload your financial statements, select industry, and download AI-enhanced reports.")

    uploaded_file = st.file_uploader("Upload PDF (Income Statement, Balance Sheet, Cash Flow)", type="pdf")
    industry = st.selectbox("Select Industry", [
        "Manufacturing", "Retail", "Technology", "Healthcare", "Energy", "Finance",
        "Construction", "Real Estate", "Telecom", "Transportation", "Hospitality", "Education"
    ])
    formats = st.multiselect("Select Report Output Format", ["PDF", "Word", "Excel"])

    if uploaded_file and st.button("Process and Analyze"):
        with st.spinner("Processing the PDF and analyzing..."):
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
                kpis = extract_kpis(text)

                if not kpis:
                    st.error("No financial metrics were extracted. Please upload a clearer PDF.")
                    return

                groq_result = groq_analyze(text, industry)

                st.success("✅ Analysis Complete")
                st.subheader("📌 KPI Summary")
                df = pd.DataFrame([
                    {
                        "Metric": k,
                        "Company": v,
                        "Benchmark": groq_result["benchmarks"].get(k, 0),
                        "Variance": v - groq_result["benchmarks"].get(k, 0),
                        "Variance %": round(((v - groq_result["benchmarks"].get(k, 0)) / groq_result["benchmarks"].get(k, 1) * 100), 2) if groq_result["benchmarks"].get(k) else 0
                    }
                    for k, v in kpis.items()
                ])
                st.dataframe(df, use_container_width=True)

                st.subheader("🧠 AI Commentary")
                st.markdown(f"**Analysis:** {groq_result['analysis']}")
                st.markdown(f"**Recommendations:** {groq_result['recommendations']}")

                st.subheader("📥 Download Reports")
                if "Excel" in formats:
                    excel_file = generate_excel(kpis, groq_result["benchmarks"])
                    st.download_button("Download Excel", excel_file, "KPI_Report.xlsx")
                if "Word" in formats:
                    word_file = generate_word(groq_result["analysis"], groq_result["recommendations"])
                    st.download_button("Download Word", word_file, "AI_Analysis_Report.docx")
                if "PDF" in formats:
                    pdf_file = generate_pdf(kpis, groq_result["benchmarks"])
                    st.download_button("Download PDF", pdf_file, "KPI_Chart_Report.pdf", mime="application/pdf")
            except Exception as e:
                st.error(f"An error occurred while processing the file: {e}")

if __name__ == "__main__":
    main()
